import io
import docx
import openpyxl
from app.parser.models import ParsedDocument, ParsedPage
from app.parser.pdf_parser import ParsingError

class OfficeParser:
    @staticmethod
    def parse(file_bytes: bytes, file_ext: str) -> ParsedDocument:
        parsed_doc = ParsedDocument()
        
        try:
            if file_ext == ".docx":
                doc = docx.Document(io.BytesIO(file_bytes))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if text:
                    parsed_doc.pages.append(ParsedPage(page_number=1, text=text))
                    
            elif file_ext == ".xlsx":
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                text_parts = []
                for sheet in wb.worksheets:
                    text_parts.append(f"--- Sheet: {sheet.title} ---")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) for cell in row if cell is not None])
                        if row_text.strip():
                            text_parts.append(row_text)
                
                text = "\n".join(text_parts)
                if text:
                    parsed_doc.pages.append(ParsedPage(page_number=1, text=text))
            
            else:
                raise ParsingError(f"Unsupported office extension: {file_ext}")
                
            if not parsed_doc.pages:
                raise ParsingError("No text extracted from document.")
                
            return parsed_doc
            
        except Exception as e:
            if isinstance(e, ParsingError):
                raise
            raise ParsingError(f"Failed to parse office document: {str(e)}")
